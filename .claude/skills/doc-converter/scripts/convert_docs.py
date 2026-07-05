"""
裁判文书格式转化工具。

将 .doc / .docx 格式的裁判文书转化为干净的纯文本（.txt）。
支持批量处理、自动格式检测、文书结构分段、元数据提取。

引擎优先级:
  .doc  → antiword (Git Bash 自带) → LibreOffice fallback → mammoth fallback
  .docx → python-docx

用法:
  python convert_docs.py --input-dir <dir> --output-dir <dir> [--index-output <path>]
"""

import os
import re
import sys
import subprocess
import shutil
import argparse
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Tuple, Optional

import chardet

# ============ 裁判文书结构标记 ============

SECTION_MARKERS = {
    "facts": [
        r"本院查明", r"经审理查明", r"再审查明", r"原审查明",
        r"一审认定", r"二审认定",
    ],
    "reasoning": [
        r"本院认为", r"本院审查认为", r"本院再审认为",
        r"本院经审查认为",
    ],
    "judgment": [
        r"判决如下", r"裁定如下", r"判决[：:]", r"裁定[：:]",
    ],
    "appeal": [
        r"如不服本判决", r"如不服本裁定",
    ],
}


def detect_format(filepath: str) -> str:
    """
    检测文件格式。

    Returns:
        'doc' | 'docx' | 'txt' | 'unknown'
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".doc":
        return "doc"
    elif ext == ".docx":
        return "docx"
    elif ext == ".txt":
        return "txt"

    # 尝试通过 magic bytes 判断
    try:
        with open(filepath, "rb") as f:
            header = f.read(8)
        if header[:4] == b"PK\x03\x04":
            return "docx"  # docx 本质是 zip
        if header[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return "doc"   # OLE2 容器
    except Exception:
        pass

    return "unknown"


def find_antiword() -> Optional[str]:
    """查找 antiword 可执行文件"""
    # 先检查 PATH
    found = shutil.which("antiword")
    if found:
        return found

    # Windows 常见安装路径
    candidates = [
        r"C:\Program Files\Git\mingw64\bin\antiword.exe",
        r"C:\Program Files\Git\usr\bin\antiword.exe",
        r"C:\cygwin64\bin\antiword.exe",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p

    return None


def find_libreoffice() -> Optional[str]:
    """查找 LibreOffice"""
    found = shutil.which("soffice")
    if found:
        return found

    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p

    return None


def detect_encoding(filepath: str) -> str:
    """检测文件编码"""
    try:
        with open(filepath, "rb") as f:
            raw = f.read(10000)
        result = chardet.detect(raw)
        return result.get("encoding", "utf-8") or "utf-8"
    except Exception:
        return "utf-8"


def convert_doc_antiword(filepath: str) -> Optional[str]:
    """
    使用 antiword 将 .doc 转为纯文本。

    Returns:
        文本内容，失败返回 None
    """
    antiword = find_antiword()
    if not antiword:
        return None

    try:
        # 不使用 text=True，Windows 上默认 GBK 会破坏 UTF-8 输出
        result = subprocess.run(
            [antiword, "-m", "UTF-8.txt", filepath],
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.decode("utf-8", errors="replace")

        # 尝试无 mapping 参数
        if result.returncode != 0 or not result.stdout.strip():
            result2 = subprocess.run(
                [antiword, filepath],
                capture_output=True,
                timeout=30,
            )
            if result2.returncode == 0 and result2.stdout.strip():
                # 尝试检测编码
                raw = result2.stdout
                detected = chardet.detect(raw)
                encoding = detected.get("encoding", "utf-8") or "utf-8"
                return raw.decode(encoding, errors="replace")
    except FileNotFoundError:
        pass
    except subprocess.TimeoutExpired:
        print(f"  [超时] {os.path.basename(filepath)}: antiword 超时")
    except Exception as e:
        print(f"  [错误] {os.path.basename(filepath)}: {e}")

    return None


def convert_doc_fallback(filepath: str) -> Optional[str]:
    """
    .doc 转化备选方案。
    尝试 LibreOffice → mammoth 的降级链。
    """
    # 方案 A: LibreOffice headless
    libreoffice = find_libreoffice()
    if libreoffice:
        try:
            import tempfile
            with tempfile.TemporaryDirectory() as tmpdir:
                subprocess.run(
                    [libreoffice, "--headless", "--convert-to", "txt:Text",
                     "--outdir", tmpdir, filepath],
                    check=True, timeout=60,
                    capture_output=True,
                )
                # 找到生成的 txt 文件
                basename = os.path.splitext(os.path.basename(filepath))[0]
                txt_path = os.path.join(tmpdir, basename + ".txt")
                if os.path.exists(txt_path):
                    with open(txt_path, "r", encoding="utf-8", errors="replace") as f:
                        return f.read()
        except Exception as e:
            print(f"  [LibreOffice fallback 失败] {e}")

    # 方案 B: mammoth (纯 Python, 从 PyPI 按需安装)
    try:
        import mammoth
        with open(filepath, "rb") as f:
            result = mammoth.extract_raw_text(f)
        if result.value.strip():
            return result.value
    except ImportError:
        pass
    except Exception as e:
        print(f"  [mammoth fallback 失败] {e}")

    return None


def convert_docx(filepath: str) -> Optional[str]:
    """
    使用 python-docx 将 .docx 转为纯文本。
    自动排除页眉页脚，保留段落结构。
    """
    try:
        from docx import Document
        doc = Document(filepath)

        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格内容（部分判决信息在表格中）
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        row_texts.append(ct)
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        return "\n".join(paragraphs)

    except ImportError:
        print("  [错误] python-docx 未安装，请运行: pip install python-docx")
    except Exception as e:
        print(f"  [错误] 解析 {os.path.basename(filepath)}: {e}")

    return None


def convert_file(filepath: str, output_dir: str) -> Tuple[Optional[str], str]:
    """
    转化单个文件。

    Args:
        filepath: 输入文件路径
        output_dir: 输出目录

    Returns:
        (output_path, status)
        status: 'success' | 'failed_antiword' | 'failed_format' | 'empty_output'
    """
    fmt = detect_format(filepath)
    basename = os.path.splitext(os.path.basename(filepath))[0]
    output_path = os.path.join(output_dir, basename + ".txt")

    text = None
    status = "success"

    if fmt == "txt":
        # 已经是纯文本，直接复制并清理
        enc = detect_encoding(filepath)
        with open(filepath, "r", encoding=enc, errors="replace") as f:
            text = f.read()
        status = "success"

    elif fmt == "docx":
        text = convert_docx(filepath)
        if text is None:
            status = "failed_format"

    elif fmt == "doc":
        text = convert_doc_antiword(filepath)
        if text is None:
            text = convert_doc_fallback(filepath)
        if text is None:
            status = "failed_antiword"

    else:
        status = "failed_format"

    # 检查输出是否为空
    if text and text.strip():
        # 基本清理
        text = clean_text(text)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)
    elif status == "success":
        status = "empty_output"
        # 仍然写一个标记文件
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"[空文件] 原文件: {filepath}\n")

    return output_path, status


def clean_text(text: str) -> str:
    """基本文本清理"""
    # 移除多余空行（保留单个换行）
    text = re.sub(r'\n{3,}', '\n\n', text)
    # 移除行首行尾空白
    text = '\n'.join(line.strip() for line in text.split('\n'))
    # 移除空行开头的内容（antiword 有时会产生）
    text = text.strip()
    return text


def segment_judgment(text: str) -> dict:
    """
    按裁判文书五段式结构分段。

    Returns:
        {
            'header': str,    # 首部（从开头到"经审理查明"之前）
            'facts': str,     # 事实部分
            'reasoning': str, # 本院认为
            'judgment': str,  # 裁判结果
            'tail': str,      # 尾部
        }
    """
    sections = {
        "header": "",
        "facts": "",
        "reasoning": "",
        "judgment": "",
        "tail": "",
    }

    # 找各段落边界
    facts_pos = None
    reasoning_pos = None
    judgment_pos = None
    appeal_pos = None

    for marker in SECTION_MARKERS["facts"]:
        m = re.search(marker, text)
        if m:
            facts_pos = min(facts_pos, m.start()) if facts_pos else m.start()

    for marker in SECTION_MARKERS["reasoning"]:
        m = re.search(marker, text)
        if m:
            reasoning_pos = min(reasoning_pos, m.start()) if reasoning_pos else m.start()

    for marker in SECTION_MARKERS["judgment"]:
        m = re.search(marker, text)
        if m:
            judgment_pos = min(judgment_pos, m.start()) if judgment_pos else m.start()

    for marker in SECTION_MARKERS["appeal"]:
        m = re.search(marker, text)
        if m:
            appeal_pos = min(appeal_pos, m.start()) if appeal_pos else m.start()

    # 按边界切分
    if facts_pos is not None:
        sections["header"] = text[:facts_pos].strip()
    else:
        sections["header"] = ""  # 无法识别

    if reasoning_pos is not None:
        start = facts_pos if facts_pos else (len(sections["header"]))
        sections["facts"] = text[start:reasoning_pos].strip()
    elif facts_pos is not None:
        sections["facts"] = text[facts_pos:].strip()

    if judgment_pos is not None and reasoning_pos is not None:
        sections["reasoning"] = text[reasoning_pos:judgment_pos].strip()
    elif reasoning_pos is not None:
        sections["reasoning"] = text[reasoning_pos:].strip()

    if judgment_pos is not None:
        end = appeal_pos if appeal_pos else len(text)
        sections["judgment"] = text[judgment_pos:end].strip()

    if appeal_pos is not None:
        sections["tail"] = text[appeal_pos:].strip()

    return sections


def extract_metadata(text: str) -> dict:
    """
    从裁判文书首部提取元数据。

    Returns:
        {
            '案号': str or None,
            '审理法院': str or None,
            '审结日期': str or None,
            '审级': str or None,
            '文书类型': str or None,
            '案由': str or None,
        }
    """
    metadata = {
        "案号": None,
        "审理法院": None,
        "审结日期": None,
        "审级": None,
        "文书类型": None,
        "案由": None,
    }

    # 只在前 2000 字符中搜索（首部）
    header = text[:2000]

    # 案号: (YYYY)XX...号
    m = re.search(r'[\(（](\d{4})[\)）]([一-龥\w]+)号', header)
    if m:
        metadata["案号"] = m.group(0)

    # 审理法院: XX人民法院
    m = re.search(r'([一-龥]{2,}(?:市|省|县|区|自治州|旗)?'
                  r'(?:高级|中级|基层)?人民法院)', header)
    if m:
        metadata["审理法院"] = m.group(1)

    # 审结日期 / 裁判日期
    m = re.search(r'(\d{4})[\.年](\d{1,2})[\.月](\d{1,2})', header)
    if m:
        metadata["审结日期"] = f"{m.group(1)}.{int(m.group(2)):02d}.{int(m.group(3)):02d}"

    # 审级
    if re.search(r'一审|初审', header):
        metadata["审级"] = "一审"
    elif re.search(r'二审|终审|上诉', header):
        metadata["审级"] = "二审"
    elif re.search(r'再审|审判监督', header):
        metadata["审级"] = "再审"

    # 文书类型
    m = re.search(r'(民事|行政|刑事|执行|赔偿)?(判决书|裁定书|决定书|调解书)', header)
    if m:
        metadata["文书类型"] = m.group(0)

    # 案由
    m = re.search(r'案由[：:]\s*(.+)', header)
    if m:
        metadata["案由"] = m.group(1).strip()

    return metadata


def batch_convert(
    input_dir: str,
    output_dir: str,
    index_output: Optional[str] = None,
    formats: Optional[list] = None,
) -> dict:
    """
    批量转化目录下的所有裁判文书。

    Args:
        input_dir: 输入目录（可递归）
        output_dir: 输出目录
        index_output: index_raw.csv 输出路径
        formats: 允许的格式列表，默认 ['doc', 'docx', 'txt']

    Returns:
        {
            'total': 总数,
            'success': 成功数,
            'failed': 失败数,
            'skipped': 跳过数,
            'records': [...],  # 每条记录的元数据
        }
    """
    if formats is None:
        formats = ["doc", "docx", "txt"]

    os.makedirs(output_dir, exist_ok=True)

    # 收集所有文件
    files = []
    for root, dirs, filenames in os.walk(input_dir):
        for fn in filenames:
            ext = os.path.splitext(fn)[1].lower()
            if ext.lstrip(".") in formats or ext in formats:
                files.append(os.path.join(root, fn))

    if not files:
        print(f"在 {input_dir} 中未找到任何 {'/'.join(formats)} 文件")
        return {"total": 0, "success": 0, "failed": 0, "skipped": 0, "records": []}

    print(f"找到 {len(files)} 个文件，开始转化...")

    results = {
        "total": len(files),
        "success": 0,
        "failed": 0,
        "skipped": 0,
        "records": [],
    }

    for i, filepath in enumerate(files, 1):
        basename = os.path.basename(filepath)
        output_path = os.path.join(output_dir, os.path.splitext(basename)[0] + ".txt")

        # 检查是否已转化（跳过已有文件）
        if os.path.exists(output_path):
            results["skipped"] += 1
            print(f"  [{i}/{len(files)}] 跳过 (已存在): {basename}")
            continue

        print(f"  [{i}/{len(files)}] 转化: {basename}...", end=" ")

        out_path, status = convert_file(filepath, output_dir)

        if status == "success":
            results["success"] += 1
            print("OK")
        else:
            results["failed"] += 1
            print(f"失败 ({status})")

        # 提取元数据
        record = {
            "原始文件名": basename,
            "原始路径": filepath,
            "转化后路径": out_path if out_path else "",
            "转化状态": status,
            "文件大小(字节)": os.path.getsize(filepath) if os.path.exists(filepath) else 0,
        }

        # 对成功的文件提取元数据
        if status == "success" and out_path and os.path.exists(out_path):
            try:
                with open(out_path, "r", encoding="utf-8", errors="replace") as f:
                    text = f.read()
                meta = extract_metadata(text)
                record.update(meta)
            except Exception:
                pass

        results["records"].append(record)

    # 输出索引 CSV
    if index_output:
        _write_index_csv(results["records"], index_output)

    # 汇报
    print(f"\n转化完成: 成功 {results['success']}, 失败 {results['failed']}, "
          f"跳过 {results['skipped']}, 总计 {results['total']}")

    return results


def _write_index_csv(records: list, output_path: str):
    """将记录写入 index_raw.csv"""
    import csv
    if not records:
        return

    # 收集所有字段名
    fieldnames = list(records[0].keys())
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(records)

    print(f"索引表已保存: {output_path} ({len(records)} 条记录)")


# ============ CLI ============

def main():
    parser = argparse.ArgumentParser(
        description="裁判文书格式转化工具 — 将 .doc/.docx 转为纯文本"
    )
    parser.add_argument("--input-dir", required=True, help="输入目录")
    parser.add_argument("--output-dir", required=True, help="输出目录（纯文本）")
    parser.add_argument("--index-output", default=None, help="index_raw.csv 输出路径")
    parser.add_argument("--formats", default="doc,docx,txt",
                        help="允许的文件格式，逗号分隔 (默认: doc,docx,txt)")
    args = parser.parse_args()

    formats = [f.strip() for f in args.formats.split(",")]

    batch_convert(
        input_dir=args.input_dir,
        output_dir=args.output_dir,
        index_output=args.index_output,
        formats=formats,
    )


if __name__ == "__main__":
    main()
